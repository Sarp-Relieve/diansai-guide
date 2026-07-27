#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
电赛技术报告 Markdown → Word (.docx) 转换脚本

全国统一排版（正文 ≤8 页意识；附录另说）：
  - 大标题（摘要等一级）: 三号宋体加粗
  - 小标题（1 / 1.1 / 1.1.1）: 四号宋体加粗
  - 正文: 小四号；中文宋体，英文/数字 Times New Roman
  - 行距: 固定值 22 磅
  - 图名图下、表名表上: 五号
  - 页上边距 ≥ 3cm；页码右下，从正文起编
  - 公式: 尽量保留为可编辑 OMML（简单 LaTeX/$...$）；复杂公式请在 Word 中用公式编辑器重录

依赖:
  pip install python-docx

用法:
  python md_to_docx.py input.md -o 技术报告.docx
  python md_to_docx.py input.md -o out.docx --title "自动行驶小车"
  python md_to_docx.py input.md -o out.docx --start-page 1

Markdown 约定（与 templates/report_outline.md 一致）:
  # 标题        → 封面大题 / 或一级（若只有一个 # 且无摘要则当题目）
  ## 摘要       → 一级标题
  ## 1 引言     → 一级标题
  ### 2.1 xxx   → 二级标题
  #### 3.1.1    → 三级标题
  ![图名](path) → 插图 + 下方图名（五号）
  | 表 |        → 表格；若前一行是「表x xxx」则作表名
  $$...$$ / $...$ → 公式段落（简易）
  ``` 代码块   → 等宽段落（建议正文少用，源码放附录）
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, Twips, RGBColor, Inches, Emu
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("缺少 python-docx，请执行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------- 字号（磅）全国统一 ----------------
SIZE_YIHAO_SAN = 16      # 三号 ≈ 16pt
SIZE_XIAOBIAO_SI = 14    # 四号 ≈ 14pt
SIZE_ZHENGWEN_XIAOSI = 12  # 小四 ≈ 12pt
SIZE_TU_BIAO_WU = 10.5   # 五号 ≈ 10.5pt
LINE_SPACING_PT = 22.0   # 固定行距 22 磅
MARGIN_TOP_CM = 3.0      # 上方 ≥3cm
MARGIN_OTHER_CM = 2.5
FONT_CN = "宋体"
FONT_EN = "Times New Roman"
FONT_CODE = "Consolas"


def set_run_font(run, size_pt: float, bold: bool = False, east_asia: str = FONT_CN, ascii_font: str = FONT_EN):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = ascii_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:cs"), ascii_font)


def set_paragraph_format(p, line_pt: float = LINE_SPACING_PT, first_indent_cm: Optional[float] = None,
                         space_before: float = 0, space_after: float = 0,
                         align=None):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent_cm is not None:
        pf.first_line_indent = Cm(first_indent_cm)
    if align is not None:
        p.alignment = align


def add_page_number(paragraph):
    """页脚页码域：右对齐当前页码"""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run2 = paragraph.add_run()
    run2._r.append(instr)
    run3 = paragraph.add_run()
    run3._r.append(fld_char_end)
    set_run_font(run, SIZE_TU_BIAO_WU)
    set_run_font(run2, SIZE_TU_BIAO_WU)
    set_run_font(run3, SIZE_TU_BIAO_WU)


def setup_document(doc: Document, start_page: int = 1):
    section = doc.sections[0]
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_OTHER_CM)
    section.left_margin = Cm(MARGIN_OTHER_CM)
    section.right_margin = Cm(MARGIN_OTHER_CM)
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)

    # 页脚页码右下
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)

    # 起始页码
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start_page))

    # 默认正文样式
    try:
        normal = doc.styles["Normal"]
        normal.font.name = FONT_EN
        normal.font.size = Pt(SIZE_ZHENGWEN_XIAOSI)
        rPr = normal.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:ascii"), FONT_EN)
        rFonts.set(qn("w:hAnsi"), FONT_EN)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(LINE_SPACING_PT)
    except Exception:
        pass


def add_heading_styled(doc: Document, text: str, level: int):
    """
    level 1: 三号加粗（摘要、1 引言…）
    level 2/3: 四号加粗
    """
    p = doc.add_paragraph()
    if level == 1:
        size, space_before, space_after = SIZE_YIHAO_SAN, 12, 6
    else:
        size, space_before, space_after = SIZE_XIAOBIAO_SI, 10, 4
    set_paragraph_format(p, space_before=space_before, space_after=space_after, first_indent_cm=0)
    run = p.add_run(text.strip())
    set_run_font(run, size, bold=True)
    return p


def add_body_paragraph(doc: Document, text: str, first_indent: bool = True):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        first_indent_cm=0.74 if first_indent else 0,  # 约 2 字符
        space_before=0,
        space_after=0,
    )
    add_mixed_runs(p, text.strip())
    return p


def add_caption(doc: Document, text: str, kind: str = "figure"):
    """图名/表名 五号居中"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_indent_cm=0, space_before=4, space_after=6,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text.strip())
    set_run_font(run, SIZE_TU_BIAO_WU, bold=False)
    return p


def add_mixed_runs(paragraph, text: str):
    """简单处理 **粗体**、`code`、普通文本"""
    if not text:
        return
    # 拆分 **bold** 与 `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _add_text_run(paragraph, text[pos:m.start()], bold=False, code=False)
        token = m.group(0)
        if token.startswith("**"):
            _add_text_run(paragraph, token[2:-2], bold=True, code=False)
        elif token.startswith("`"):
            _add_text_run(paragraph, token[1:-1], bold=False, code=True)
        elif token.startswith("*"):
            _add_text_run(paragraph, token[1:-1], bold=True, code=False)
        pos = m.end()
    if pos < len(text):
        _add_text_run(paragraph, text[pos:], bold=False, code=False)


def _add_text_run(paragraph, text: str, bold: bool, code: bool):
    if not text:
        return
    run = paragraph.add_run(text)
    if code:
        set_run_font(run, SIZE_ZHENGWEN_XIAOSI - 1, bold=False, east_asia=FONT_CODE, ascii_font=FONT_CODE)
    else:
        set_run_font(run, SIZE_ZHENGWEN_XIAOSI, bold=bold)


def add_formula_paragraph(doc: Document, latex: str):
    """
    简易公式：写入居中段落，外加 $ 提示。
    复杂公式请在 Word 公式编辑器中重录（竞赛要求）。
    """
    latex = latex.strip().strip("$").strip()
    p = doc.add_paragraph()
    set_paragraph_format(p, first_indent_cm=0, space_before=6, space_after=6,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(latex)
    set_run_font(run, SIZE_ZHENGWEN_XIAOSI, bold=False, east_asia=FONT_EN, ascii_font=FONT_EN)
    # 批注式提示
    tip = doc.add_paragraph()
    set_paragraph_format(tip, first_indent_cm=0, space_before=0, space_after=4,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    r2 = tip.add_run("【请用 Word 公式编辑器按上式重录】")
    set_run_font(r2, SIZE_TU_BIAO_WU, bold=False)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p


def add_code_block(doc: Document, code: str):
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_indent_cm=0, line_pt=16, space_before=0, space_after=0)
        run = p.add_run(line if line.strip() else " ")
        set_run_font(run, 9, bold=False, east_asia=FONT_CODE, ascii_font=FONT_CODE)


def add_image(doc: Document, path: str, max_width_cm: float = 14.0):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_indent_cm=0, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=6, space_after=2)
    if not os.path.isfile(path):
        run = p.add_run(f"【缺图: {path}】")
        set_run_font(run, SIZE_TU_BIAO_WU)
        return p
    run = p.add_run()
    try:
        run.add_picture(path, width=Cm(max_width_cm))
    except Exception as e:
        run = p.add_run(f"【插图失败: {path} ({e})】")
        set_run_font(run, SIZE_TU_BIAO_WU)
    return p


def parse_table_block(lines: List[str], start: int) -> Tuple[List[List[str]], int]:
    """从 start 解析 markdown 表格，返回 rows, next_index"""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip().startswith("|"):
            break
        # 分隔行 |---|
        if re.match(r"^\s*\|?\s*:?-{3,}", line.replace(" ", "")) or re.match(
            r"^\s*\|(\s*:?-+:?\s*\|)+\s*$", line
        ):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: List[List[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, first_indent_cm=0, line_pt=18, space_before=2, space_after=2)
            run = p.add_run(text)
            set_run_font(run, SIZE_TU_BIAO_WU, bold=(ri == 0))
    # 表后空行
    doc.add_paragraph()


def is_caption_line(line: str) -> Optional[str]:
    s = line.strip()
    if re.match(r"^图\s*\d+", s) or re.match(r"^表\s*\d+", s):
        return s
    if re.match(r"^Figure\s*\d+", s, re.I) or re.match(r"^Table\s*\d+", s, re.I):
        return s
    return None


def convert_md_to_docx(md_text: str, out_path: str, title: Optional[str] = None, start_page: int = 1,
                       base_dir: Optional[str] = None):
    base_dir = base_dir or os.getcwd()
    doc = Document()
    setup_document(doc, start_page=start_page)

    lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    n = len(lines)
    pending_table_caption: Optional[str] = None
    in_appendix = False

    # 可选封面题
    if title:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_indent_cm=0, space_before=24, space_after=24,
                             align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(title)
        set_run_font(r, 18, bold=True)

    while i < n:
        line = lines[i]
        raw = line.rstrip("\n")
        stripped = raw.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i < n:
                i += 1
            add_code_block(doc, "\n".join(buf))
            continue

        # 块公式 $$
        if stripped.startswith("$$"):
            if stripped.endswith("$$") and len(stripped) > 4:
                add_formula_paragraph(doc, stripped)
                i += 1
                continue
            buf = [stripped.lstrip("$")]
            i += 1
            while i < n and "$$" not in lines[i]:
                buf.append(lines[i])
                i += 1
            if i < n:
                buf.append(lines[i].replace("$$", ""))
                i += 1
            add_formula_paragraph(doc, "\n".join(buf))
            continue

        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # 去掉末尾手动编号重复
            if re.search(r"附录", text):
                in_appendix = True
            # # 仅作文档题时：若已有 title 参数则跳过重复
            if level == 1 and title and text == title:
                i += 1
                continue
            # markdown # → 一级（三号），## → 一级或二级：
            # 约定：## 摘要 / ## 1 引言 为一级；### 为二级；#### 为三级
            if level == 1:
                add_heading_styled(doc, text, 1)
            elif level == 2:
                add_heading_styled(doc, text, 1)
            elif level == 3:
                add_heading_styled(doc, text, 2)
            else:
                add_heading_styled(doc, text, 2)
            i += 1
            continue

        # 图片 ![alt](path)
        m = re.match(r"^!\[(.*?)\]\((.+?)\)\s*$", stripped)
        if m:
            alt, path = m.group(1), m.group(2).strip().strip('"')
            if not os.path.isabs(path):
                path = os.path.join(base_dir, path)
            add_image(doc, path)
            cap = alt.strip() or pending_table_caption
            # 下一行若是图名则优先
            if i + 1 < n and is_caption_line(lines[i + 1]):
                i += 1
                add_caption(doc, lines[i].strip(), "figure")
            elif cap:
                # alt 当图名
                if not re.match(r"^图", cap):
                    cap = cap
                add_caption(doc, cap, "figure")
            i += 1
            continue

        # 表名行（表在下一行）
        cap = is_caption_line(stripped)
        if cap and cap.startswith("表"):
            pending_table_caption = cap
            i += 1
            # 跳过空行
            while i < n and not lines[i].strip():
                i += 1
            if i < n and lines[i].strip().startswith("|"):
                rows, i = parse_table_block(lines, i)
                add_caption(doc, pending_table_caption, "table")
                add_table(doc, rows)
                pending_table_caption = None
            else:
                add_caption(doc, cap, "table")
            continue

        # 图名独占行（无图时也保留）
        if cap and cap.startswith("图"):
            add_caption(doc, cap, "figure")
            i += 1
            continue

        # 表格
        if stripped.startswith("|"):
            rows, i = parse_table_block(lines, i)
            if pending_table_caption:
                add_caption(doc, pending_table_caption, "table")
                pending_table_caption = None
            add_table(doc, rows)
            continue

        # 行内公式段落（整行 $...$）
        if stripped.startswith("$") and stripped.endswith("$") and stripped.count("$") == 2:
            add_formula_paragraph(doc, stripped)
            i += 1
            continue

        # 无序列表 / 有序列表 → 正文缩进
        mlist = re.match(r"^([-*+]|\d+[\.、）)])\s+(.*)$", stripped)
        if mlist:
            p = doc.add_paragraph()
            set_paragraph_format(p, first_indent_cm=0, space_before=0, space_after=0)
            # 悬挂一点
            p.paragraph_format.left_indent = Cm(0.74)
            add_mixed_runs(p, stripped)
            i += 1
            continue

        # 普通正文（合并续行到空行或下一结构）
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            ns = nxt.strip()
            if not ns:
                break
            if ns.startswith("#") or ns.startswith("```") or ns.startswith("|") \
                    or ns.startswith("$$") or ns.startswith("![") \
                    or is_caption_line(ns) \
                    or re.match(r"^([-*+]|\d+[\.、）)])\s+", ns):
                break
            # 软换行合并
            buf.append(ns)
            i += 1
        text = "".join(buf) if all(len(x) < 40 for x in buf) else "".join(buf)
        # 中文报告习惯：段落内直接拼接（markdown 单换行常为同一段）
        text = re.sub(r"\s+", " ", " ".join(buf)).strip()
        # 还原中文无空格：若多为中文则去空格间多余
        if len(re.findall(r"[一-鿿]", text)) > len(text) * 0.3:
            # 仅去掉「中文 中文」之间的空格，保留英文单词空格
            text = re.sub(r"([一-鿿])\s+([一-鿿])", r"\1\2", text)
            text = re.sub(r"([一-鿿])\s+([，。；：、）】》])", r"\1\2", text)
            text = re.sub(r"([（【《])\s+([一-鿿])", r"\1\2", text)
        add_body_paragraph(doc, text, first_indent=True)

    doc.save(out_path)
    return out_path


def main(argv=None):
    ap = argparse.ArgumentParser(description="电赛技术报告 Markdown → docx（全国统一排版）")
    ap.add_argument("input", help="输入 Markdown 文件")
    ap.add_argument("-o", "--output", help="输出 docx 路径", default=None)
    ap.add_argument("--title", help="封面/文首题目", default=None)
    ap.add_argument("--start-page", type=int, default=1, help="页码起始，默认 1（正文起编）")
    args = ap.parse_args(argv)

    in_path = Path(args.input).expanduser().resolve()
    if not in_path.is_file():
        print(f"找不到输入文件: {in_path}", file=sys.stderr)
        return 2

    out = args.output
    if not out:
        out = str(in_path.with_suffix(".docx"))
    out_path = str(Path(out).expanduser().resolve())

    md_text = in_path.read_text(encoding="utf-8")
    convert_md_to_docx(
        md_text,
        out_path,
        title=args.title,
        start_page=args.start_page,
        base_dir=str(in_path.parent),
    )
    print(f"已生成: {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main() or 0)
